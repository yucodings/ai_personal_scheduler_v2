from io import BytesIO
from docx import Document


def parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content)); lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, 1):
        lines.append(f"[Table {table_index}]")
        lines.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
    return "\n".join(lines)

